from __future__ import annotations

import hashlib
import ipaddress
import json
import mimetypes
import re
import socket
import uuid
import urllib.request
import urllib.parse
from dataclasses import dataclass, field, replace
from datetime import datetime, timezone
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any, Callable, Dict, Iterable, List, Optional, Protocol, Tuple

from aragbiz.model_farm import ModelCallContext, ModelFarmError, ModelFarmService, ModelGateway


CUSTOM_CHUNKING_EXTENSIONS = {".aratxt", ".arajson", ".aramd"}
CHUNKING_STRATEGIES = {
    "fixed_size",
    "sliding_window_overlap",
    "header_based",
    "semantic",
    "recursive",
    "structure_aware_recursive",
    "hierarchical_parent_child",
    "structure_aware_custom",
}
WIXQA_MINILM_PROFILE_ID = "wixqa_minilm_structure_v1"
STRUCTURE_SEPARATOR_TYPES = {
    "heading",
    "subheading",
    "paragraph",
    "list",
    "sentence",
    "token",
}
DEFAULT_EMBEDDING_OPTIONS = {
    "hard_max_wordpieces": 512,
}
WIXQA_MINILM_EMBEDDING_OPTIONS = {
    "hard_max_wordpieces": 240,
}
WIXQA_MINILM_CHUNKING_OPTIONS = {
    "parent_document": "article_id",
    "target_body_tokens": 180,
    "soft_max_body_tokens": 210,
    "overlap_tokens": 30,
    "minimum_chunk_tokens": 60,
    "separators": ["heading", "subheading", "paragraph", "list", "sentence", "token"],
    "metadata_prefix": {
        "include_title": True,
        "include_heading_path": True,
        "include_article_type": False,
        "maximum_tokens": 40,
    },
    "rules": {
        "preserve_numbered_lists": True,
        "preserve_bullet_lists": True,
        "merge_small_adjacent_chunks": True,
        "overlap_only_within_same_section": True,
        "never_merge_across_articles": True,
    },
}
DEFAULT_KNOWLEDGE_BASE_CONFIGURATION = {
    "chunking_strategy": "sliding_window_overlap",
    "chunking_profile_id": "",
    "chunk_size": 800,
    "chunk_overlap": 120,
    "embedding_options": DEFAULT_EMBEDDING_OPTIONS,
    "chunking_options": WIXQA_MINILM_CHUNKING_OPTIONS,
    "embedding_provider": "Local",
    "embedding_model": "hash-embedding-384",
    "embedding_deployment_id": "model-local-hash-384",
    "external_processing_allowed": False,
}
LOCAL_EMBEDDING_PROVIDER = "Local"
HASH_EMBEDDING_MODEL = "hash-embedding-384"
SENTENCE_TRANSFORMER_MINILM_MODEL = "sentence-transformers/all-MiniLM-L6-v2"
SUPPORTED_LOCAL_EMBEDDING_MODELS = {
    HASH_EMBEDDING_MODEL,
    SENTENCE_TRANSFORMER_MINILM_MODEL,
}
SUPPORTED_FILE_EXTENSIONS = {
    ".txt",
    ".json",
    ".jsonl",
    ".md",
    ".pdf",
    ".docx",
    *CUSTOM_CHUNKING_EXTENSIONS,
}
WIXQA_CORPUS_ID = "wix_kb_corpus"
WIXQA_CORPUS_EXPECTED_DOCUMENTS = 6221
WIXQA_CORPUS_URI = "dataset://Wix/WixQA/wix_kb_corpus"


class KnowledgeProcessingError(ValueError):
    """Raised when a knowledge source cannot be loaded or processed."""


class KnowledgeImportCancelled(KnowledgeProcessingError):
    """Raised when a durable knowledge import is cancelled."""


@dataclass(frozen=True)
class KnowledgeBaseRecord:
    id: str
    name: str
    description: str = ""
    status: str = "empty"
    document_count: int = 0
    chunk_count: int = 0
    embedding_model: str = ""
    created_at: str = ""
    updated_at: str = ""
    metadata: Dict[str, Any] = field(default_factory=dict)
    error: Optional[str] = None


@dataclass(frozen=True)
class KnowledgeIndexVersionRecord:
    id: str
    knowledge_base_id: str
    status: str
    chunking_configuration: Dict[str, Any] = field(default_factory=dict)
    embedding_deployment_id: str = ""
    embedding_provider: str = ""
    embedding_model: str = ""
    embedding_dimension: int = 0
    document_count: int = 0
    chunk_count: int = 0
    error: str = ""
    created_at: str = ""
    activated_at: str = ""


@dataclass(frozen=True)
class DataSourceRecord:
    id: str
    knowledge_base_id: str
    source_type: str
    uri: str
    status: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class KnowledgeDocumentInput:
    title: str
    text: str
    source_type: str
    uri: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class StoredKnowledgeDocument:
    id: str
    knowledge_base_id: str
    source_id: str
    title: str
    content_hash: str
    text: str
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True)
class StoredKnowledgeChunk:
    id: str
    knowledge_base_id: str
    document_id: str
    chunk_index: int
    text: str
    token_count: int
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding_model: str = ""
    embedding_dimension: int = 0
    has_embedding: bool = False
    index_version_id: str = ""
    parent_chunk_id: str = ""


@dataclass(frozen=True)
class ProcessingTraceStep:
    step: str
    status: str
    detail: str
    metadata: Dict[str, Any] = field(default_factory=dict)
    started_at: str = ""
    finished_at: str = ""


@dataclass(frozen=True)
class IngestionSummary:
    knowledge_base_id: str
    source_id: Optional[str]
    status: str
    documents_added: int
    documents_skipped: int
    chunks_added: int
    error: Optional[str] = None


class KnowledgeRepository(Protocol):
    def initialize(self) -> None:
        """Create storage tables or files if needed."""

    def create_knowledge_base(
        self,
        name: str,
        description: str = "",
        metadata: Optional[Dict[str, Any]] = None,
    ) -> KnowledgeBaseRecord:
        """Create and return a knowledge base."""

    def list_knowledge_bases(self) -> List[KnowledgeBaseRecord]:
        """Return knowledge bases with aggregate counts."""

    def get_knowledge_base(self, knowledge_base_id: str) -> KnowledgeBaseRecord:
        """Return one knowledge base."""

    def update_knowledge_base(
        self,
        knowledge_base_id: str,
        *,
        name: Optional[str] = None,
        description: Optional[str] = None,
        status: Optional[str] = None,
        embedding_model: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        error: Optional[str] = None,
    ) -> None:
        """Update knowledge-base status fields."""

    def delete_knowledge_base(self, knowledge_base_id: str) -> None:
        """Delete a knowledge base and all child records."""

    def create_data_source(
        self,
        knowledge_base_id: str,
        source_type: str,
        uri: str,
        status: str,
        metadata: Dict[str, Any],
    ) -> DataSourceRecord:
        """Create a data-source record."""

    def update_data_source(
        self,
        source_id: str,
        *,
        status: str,
        metadata: Dict[str, Any],
    ) -> None:
        """Update data-source processing status and metadata."""

    def existing_hashes(self, knowledge_base_id: str) -> set[str]:
        """Return known document hashes for deduplication."""

    def add_document(self, document: StoredKnowledgeDocument) -> None:
        """Persist a normalized document."""

    def add_documents(self, documents: List[StoredKnowledgeDocument]) -> None:
        """Persist normalized documents in one repository operation."""

    def list_documents(self, knowledge_base_id: str) -> List[StoredKnowledgeDocument]:
        """Return documents for a knowledge base."""

    def list_documents_page(
        self,
        knowledge_base_id: str,
        *,
        query: str = "",
        offset: int = 0,
        limit: int = 25,
    ) -> Tuple[List[StoredKnowledgeDocument], int]:
        """Return one searchable document page and the total match count."""

    def list_wixqa_source_record_ids(self, knowledge_base_id: str) -> List[str]:
        """Return imported WixQA corpus record IDs without loading document text."""

    def get_document(self, knowledge_base_id: str, document_id: str) -> StoredKnowledgeDocument:
        """Return one document in a knowledge base."""

    def update_document(self, document: StoredKnowledgeDocument) -> None:
        """Update a document record."""

    def delete_document(self, knowledge_base_id: str, document_id: str) -> None:
        """Delete a document and its chunks/embeddings."""

    def replace_chunks(self, knowledge_base_id: str, chunks: List[StoredKnowledgeChunk], embeddings: List[List[float]], model: str) -> None:
        """Replace all chunks and embeddings for a knowledge base."""

    def replace_document_chunks(
        self,
        knowledge_base_id: str,
        document_id: str,
        chunks: List[StoredKnowledgeChunk],
        embeddings: List[List[float]],
        model: str,
    ) -> None:
        """Replace chunks and embeddings for one document."""

    def append_chunks(self, chunks: List[StoredKnowledgeChunk], embeddings: List[List[float]], model: str) -> None:
        """Append chunks and embeddings."""

    def list_chunks(self, knowledge_base_id: str, limit: int = 100) -> List[StoredKnowledgeChunk]:
        """Return chunks for a knowledge base."""

    def list_document_chunks(self, knowledge_base_id: str, document_id: str) -> List[StoredKnowledgeChunk]:
        """Return active chunks for one document."""

    def list_active_chunks(
        self,
        knowledge_base_id: str,
        document_ids: Optional[List[str]] = None,
    ) -> List[StoredKnowledgeChunk]:
        """Return every active chunk, optionally restricted to documents."""

    def search_chunks_by_embedding(
        self,
        knowledge_base_id: str,
        embedding: List[float],
        limit: int = 10,
    ) -> List[tuple[StoredKnowledgeChunk, float]]:
        """Return chunks ranked by vector similarity."""

    def list_ingestion_runs(self, knowledge_base_id: str, limit: int = 10) -> List[Dict[str, Any]]:
        """Return recent ingestion runs for a knowledge base."""

    def record_ingestion_run(
        self,
        knowledge_base_id: str,
        status: str,
        counts: Dict[str, int],
        error: Optional[str] = None,
        source_id: Optional[str] = None,
    ) -> None:
        """Persist an ingestion run summary."""

    def create_index_version(
        self,
        knowledge_base_id: str,
        configuration: Dict[str, Any],
        *,
        embedding_provider: str,
        embedding_model: str,
        embedding_deployment_id: str,
        embedding_dimension: int,
    ) -> KnowledgeIndexVersionRecord:
        """Create a draft immutable index version."""

    def activate_index_version(self, knowledge_base_id: str, version_id: str) -> KnowledgeIndexVersionRecord:
        """Atomically activate a completed index version."""

    def fail_index_version(self, knowledge_base_id: str, version_id: str, error: str) -> None:
        """Mark a draft version failed without replacing the active index."""

    def list_index_versions(self, knowledge_base_id: str) -> List[KnowledgeIndexVersionRecord]:
        """List index versions newest first."""


class KnowledgeService:
    def __init__(
        self,
        repository: KnowledgeRepository,
        chunker: "OverlapChunker",
        embedder: "EmbeddingModel",
        model_farm_service: Optional[ModelFarmService] = None,
        model_gateway: Optional[ModelGateway] = None,
        prepared_corpus_path: str = "data/processed/wix_kb_corpus_documents.jsonl",
        prepared_corpus_expected_documents: int = WIXQA_CORPUS_EXPECTED_DOCUMENTS,
        embedding_batch_size: int = 64,
    ):
        self.repository = repository
        self.repository.initialize()
        self.chunker = chunker
        self.default_embedder = embedder
        self._embedding_dimension = embedder.dimension
        self._embedder_cache: Dict[tuple[str, str, int], "EmbeddingModel"] = {}
        self.model_farm_service = model_farm_service
        self.model_gateway = model_gateway
        self.prepared_corpus_path = Path(prepared_corpus_path)
        self.prepared_corpus_expected_documents = max(1, int(prepared_corpus_expected_documents))
        self.embedding_batch_size = max(1, min(int(embedding_batch_size), 256))
        self._prepared_catalog_cache: Optional[Tuple[int, int, Dict[str, Any]]] = None

    def create_knowledge_base(
        self,
        name: str,
        description: str = "",
        configuration: Optional[Dict[str, Any]] = None,
    ) -> KnowledgeBaseRecord:
        self.repository.initialize()
        normalized_configuration = self._validate_configuration(configuration)
        return self.repository.create_knowledge_base(
            name=name,
            description=description,
            metadata={"configuration": normalized_configuration},
        )

    def list_knowledge_bases(self) -> List[KnowledgeBaseRecord]:
        self.repository.initialize()
        return self.repository.list_knowledge_bases()

    def get_knowledge_base(self, knowledge_base_id: str) -> KnowledgeBaseRecord:
        self.repository.initialize()
        return self.repository.get_knowledge_base(knowledge_base_id)

    def update_knowledge_base_details(
        self,
        knowledge_base_id: str,
        name: str,
        description: str = "",
        configuration: Optional[Dict[str, Any]] = None,
    ) -> KnowledgeBaseRecord:
        self.repository.initialize()
        current = self.repository.get_knowledge_base(knowledge_base_id)
        metadata = dict(current.metadata)
        if configuration is not None:
            metadata["configuration"] = self._validate_configuration(configuration)
        self.repository.update_knowledge_base(
            knowledge_base_id,
            name=name.strip(),
            description=description,
            metadata=metadata,
            error=None,
        )
        return self.repository.get_knowledge_base(knowledge_base_id)

    def delete_knowledge_base(self, knowledge_base_id: str) -> None:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        self.repository.delete_knowledge_base(knowledge_base_id)

    def list_documents(self, knowledge_base_id: str) -> List[StoredKnowledgeDocument]:
        self.repository.initialize()
        return self.repository.list_documents(knowledge_base_id)

    def list_documents_page(
        self,
        knowledge_base_id: str,
        *,
        query: str = "",
        offset: int = 0,
        limit: int = 25,
    ) -> Tuple[List[StoredKnowledgeDocument], int]:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        return self.repository.list_documents_page(
            knowledge_base_id,
            query=query,
            offset=max(0, int(offset)),
            limit=max(1, min(int(limit), 100)),
        )

    def list_wixqa_source_record_ids(self, knowledge_base_id: str) -> List[str]:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        return self.repository.list_wixqa_source_record_ids(knowledge_base_id)

    def get_document(self, knowledge_base_id: str, document_id: str) -> StoredKnowledgeDocument:
        self.repository.initialize()
        return self.repository.get_document(knowledge_base_id, document_id)

    def list_chunks(self, knowledge_base_id: str, limit: int = 100) -> List[StoredKnowledgeChunk]:
        self.repository.initialize()
        return self.repository.list_chunks(knowledge_base_id, limit=limit)

    def list_document_chunks(self, knowledge_base_id: str, document_id: str) -> List[StoredKnowledgeChunk]:
        self.repository.initialize()
        self.repository.get_document(knowledge_base_id, document_id)
        return self.repository.list_document_chunks(knowledge_base_id, document_id)

    def list_active_chunks(
        self,
        knowledge_base_id: str,
        document_ids: Optional[List[str]] = None,
    ) -> List[StoredKnowledgeChunk]:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        return self.repository.list_active_chunks(knowledge_base_id, document_ids=document_ids)

    def prepared_source_catalog(self) -> Dict[str, Any]:
        path = self.prepared_corpus_path
        unavailable = {
            "id": WIXQA_CORPUS_ID,
            "name": "WixQA knowledge corpus",
            "source": "Wix/WixQA",
            "expected_document_count": self.prepared_corpus_expected_documents,
            "available": False,
            "path": str(path),
            "download_command": "python scripts/download_wixqa.py --subset wixqa_expertwritten",
        }
        if not path.is_file():
            return {**unavailable, "error": "Prepared WixQA corpus file is not available."}
        stat = path.stat()
        cache_key = (stat.st_mtime_ns, stat.st_size)
        if self._prepared_catalog_cache and self._prepared_catalog_cache[:2] == cache_key:
            return dict(self._prepared_catalog_cache[2])
        try:
            documents, checksum = load_prepared_wixqa_corpus(
                path,
                expected_documents=self.prepared_corpus_expected_documents,
            )
            catalog = {
                **unavailable,
                "available": True,
                "document_count": len(documents),
                "character_count": sum(len(document.text) for document in documents),
                "size_bytes": stat.st_size,
                "sha256": checksum,
                "error": "",
            }
        except KnowledgeProcessingError as exc:
            catalog = {**unavailable, "size_bytes": stat.st_size, "error": str(exc)}
        self._prepared_catalog_cache = (*cache_key, catalog)
        return dict(catalog)

    def embedding_requires_remote_confirmation(self, knowledge_base_id: str) -> bool:
        configuration = self._knowledge_base_configuration(knowledge_base_id)
        deployment_id = str(configuration.get("embedding_deployment_id") or "")
        if not deployment_id or self.model_farm_service is None:
            return str(configuration.get("embedding_provider") or "").lower() != "local"
        try:
            return not self.model_farm_service.resolve(deployment_id, "embedding").is_local
        except (KeyError, ModelFarmError) as exc:
            raise KnowledgeProcessingError(str(exc)) from exc

    def embed_query(self, knowledge_base_id: str, query: str) -> tuple[List[float], str]:
        self.repository.initialize()
        configuration = self._active_index_configuration(knowledge_base_id)
        embedder = self._embedder_for_configuration(configuration)
        embedding_query = query
        if configuration.get("chunking_strategy") == "structure_aware_recursive":
            tokenizer = _load_wordpiece_tokenizer(configuration["embedding_model"])
            hard_max = int(configuration["embedding_options"]["hard_max_wordpieces"])
            embedding_query = _truncate_wordpieces(
                tokenizer,
                query,
                max(hard_max - _special_token_count(tokenizer), 1),
            )
        embeddings = embedder.embed([embedding_query])
        return (embeddings[0] if embeddings else [], embedder.model_name)

    def query_embedding_details(self, knowledge_base_id: str) -> Dict[str, Any]:
        self.repository.initialize()
        knowledge_base = self.repository.get_knowledge_base(knowledge_base_id)
        configuration = self._knowledge_base_configuration(knowledge_base_id)
        chunks = self.repository.list_chunks(knowledge_base_id, limit=1)
        chunk = chunks[0] if chunks else None
        versions = self.list_index_versions(knowledge_base_id)
        active_version = next((version for version in versions if version.status == "active"), None)
        return {
            "deployment_id": (
                active_version.embedding_deployment_id
                if active_version
                else configuration.get("embedding_deployment_id", "")
            ),
            "provider": (
                active_version.embedding_provider
                if active_version
                else configuration.get("embedding_provider", "")
            ),
            "model": (
                active_version.embedding_model
                if active_version
                else (chunk.embedding_model if chunk else knowledge_base.embedding_model)
            ),
            "dimension": (
                active_version.embedding_dimension
                if active_version
                else (chunk.embedding_dimension if chunk else self._embedding_dimension)
            ),
            "active_index_version_id": (
                active_version.id
                if active_version
                else (chunk.index_version_id if chunk else "")
            ),
            "source": "active_knowledge_index",
            "chunking_profile_id": (
                active_version.chunking_configuration.get("chunking_profile_id", "")
                if active_version
                else configuration.get("chunking_profile_id", "")
            ),
            "hard_max_wordpieces": (
                active_version.chunking_configuration.get("embedding_options", {}).get(
                    "hard_max_wordpieces",
                    512,
                )
                if active_version
                else configuration.get("embedding_options", {}).get("hard_max_wordpieces", 512)
            ),
        }

    def list_index_versions(self, knowledge_base_id: str) -> List[KnowledgeIndexVersionRecord]:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        method = getattr(self.repository, "list_index_versions", None)
        return list(method(knowledge_base_id)) if callable(method) else []

    def processing_trace(self, knowledge_base_id: str) -> List[ProcessingTraceStep]:
        self.repository.initialize()
        knowledge_base = self.repository.get_knowledge_base(knowledge_base_id)
        configuration = normalize_knowledge_base_configuration(knowledge_base.metadata.get("configuration"))
        documents = self.repository.list_documents(knowledge_base_id)
        chunks = self.repository.list_active_chunks(knowledge_base_id)
        runs = self.repository.list_ingestion_runs(knowledge_base_id, limit=8)
        embedded_chunks = sum(1 for chunk in chunks if chunk.has_embedding)
        chunk_embedding_models = sorted({chunk.embedding_model for chunk in chunks if chunk.embedding_model})
        embedding_dimension = next((chunk.embedding_dimension for chunk in chunks if chunk.embedding_dimension), self._embedding_dimension)
        source_types = sorted({document.metadata.get("source_type", "unknown") for document in documents})
        latest_run = runs[0] if runs else {}
        return [
            ProcessingTraceStep(
                step="Knowledge base selected",
                status=knowledge_base.status,
                detail=f"{knowledge_base.name} is active for document management.",
                metadata={
                    "knowledge_base_id": knowledge_base.id,
                    "document_count": knowledge_base.document_count,
                    "chunk_count": knowledge_base.chunk_count,
                },
                started_at=knowledge_base.created_at,
                finished_at=knowledge_base.updated_at,
            ),
            ProcessingTraceStep(
                step="Data source loading",
                status="completed" if documents else "waiting",
                detail=f"Loaded {len(documents)} document(s) from {', '.join(source_types) if source_types else 'no sources yet'}.",
                metadata={
                    "source_types": source_types,
                    "recent_runs": runs,
                },
                started_at=latest_run.get("started_at", ""),
                finished_at=latest_run.get("finished_at", ""),
            ),
            ProcessingTraceStep(
                step="Metadata extraction and deduplication",
                status="completed" if documents else "waiting",
                detail="Each document has title, source metadata, content hash and deduplication status.",
                metadata={
                    "content_hash_count": len(documents),
                    "content_hash_samples": [document.content_hash for document in documents[:20]],
                    "deduplication_key": "knowledge_base_id + content_hash",
                },
            ),
            ProcessingTraceStep(
                step="Chunking",
                status="completed" if chunks else "waiting",
                detail=f"Created {len(chunks)} ordered overlapping chunk(s).",
                metadata={
                    "chunk_size": configuration["chunk_size"],
                    "chunk_overlap": configuration["chunk_overlap"],
                    "configured_chunk_size": configuration["chunk_size"],
                    "configured_chunk_overlap": configuration["chunk_overlap"],
                    "chunking_strategy": configuration["chunking_strategy"],
                    "chunk_ids": [chunk.id for chunk in chunks[:20]],
                },
            ),
            ProcessingTraceStep(
                step="Embedding",
                status="completed" if embedded_chunks == len(chunks) and chunks else "waiting",
                detail=f"Embedded {embedded_chunks} of {len(chunks)} chunk(s).",
                metadata={
                    "embedding_model": knowledge_base.embedding_model or (chunk_embedding_models[0] if len(chunk_embedding_models) == 1 else ""),
                    "chunk_embedding_models": chunk_embedding_models,
                    "configured_embedding_provider": configuration["embedding_provider"],
                    "configured_embedding_model": configuration["embedding_model"],
                    "embedding_dimension": embedding_dimension,
                    "embedded_chunks": embedded_chunks,
                },
            ),
            ProcessingTraceStep(
                step="Storage",
                status="completed" if chunks else knowledge_base.status,
                detail="Documents are stored in the relational layer; chunk vectors are stored in pgVector or the JSON fallback store.",
                metadata={
                    "repository": type(self.repository).__name__,
                    "latest_run": latest_run,
                },
            ),
        ]

    def create_document(
        self,
        knowledge_base_id: str,
        title: str,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> StoredKnowledgeDocument:
        self.repository.initialize()
        configuration = self._knowledge_base_configuration(knowledge_base_id)
        embedder = self._embedder_for_configuration(configuration)
        source = self.repository.create_data_source(
            knowledge_base_id=knowledge_base_id,
            source_type="manual",
            uri=title,
            status="ready",
            metadata={"created_at": utc_now(), "entrypoint": "document_editor"},
        )
        clean = clean_text(text)
        if not clean:
            raise KnowledgeProcessingError("Document text cannot be empty.")
        digest = content_hash(clean)
        if digest in self.repository.existing_hashes(knowledge_base_id):
            raise KnowledgeProcessingError("A document with the same content already exists in this knowledge base.")
        document = StoredKnowledgeDocument(
            id=f"doc-{uuid.uuid4().hex}",
            knowledge_base_id=knowledge_base_id,
            source_id=source.id,
            title=title.strip() or "Untitled document",
            content_hash=digest,
            text=clean,
            metadata={
                **(metadata or {}),
                **_processing_metadata(configuration),
                "source_type": "manual",
                "uri": title,
                "content_hash": digest,
                "created_at": utc_now(),
                "updated_at": utc_now(),
            },
        )
        self.repository.add_document(document)
        self.reindex(knowledge_base_id)
        return document

    def update_document(
        self,
        knowledge_base_id: str,
        document_id: str,
        title: str,
        text: str,
        metadata: Optional[Dict[str, Any]] = None,
    ) -> StoredKnowledgeDocument:
        self.repository.initialize()
        existing = self.repository.get_document(knowledge_base_id, document_id)
        configuration = self._knowledge_base_configuration(knowledge_base_id)
        clean = clean_text(text)
        if not clean:
            raise KnowledgeProcessingError("Document text cannot be empty.")
        digest = content_hash(clean)
        for document in self.repository.list_documents(knowledge_base_id):
            if document.id != document_id and document.content_hash == digest:
                raise KnowledgeProcessingError("A document with the same content already exists in this knowledge base.")
        updated = StoredKnowledgeDocument(
            id=existing.id,
            knowledge_base_id=existing.knowledge_base_id,
            source_id=existing.source_id,
            title=title.strip() or existing.title,
            content_hash=digest,
            text=clean,
            metadata={
                **existing.metadata,
                **(metadata or {}),
                **_processing_metadata(configuration),
                "content_hash": digest,
                "updated_at": utc_now(),
            },
        )
        self.repository.update_document(updated)
        self.reindex(knowledge_base_id)
        return updated

    def delete_document(self, knowledge_base_id: str, document_id: str) -> None:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        self.repository.get_document(knowledge_base_id, document_id)
        self.repository.delete_document(knowledge_base_id, document_id)
        self.reindex(knowledge_base_id)

    def ingest_uploaded_file(self, knowledge_base_id: str, filename: str, content: bytes) -> IngestionSummary:
        self.repository.initialize()
        documents = load_file_documents(filename, content)
        return self._ingest_documents(
            knowledge_base_id,
            source_type="upload",
            uri=filename,
            documents=documents,
            source_metadata={"filename": filename, "size_bytes": len(content)},
        )

    def ingest_website(self, knowledge_base_id: str, url: str) -> IngestionSummary:
        self.repository.initialize()
        documents = [load_public_website(url)]
        return self._ingest_documents(
            knowledge_base_id,
            source_type="website",
            uri=url,
            documents=documents,
            source_metadata={"url": url},
        )

    def ingest_prepared_wixqa_corpus(
        self,
        knowledge_base_id: str,
        *,
        document_limit: Optional[int] = None,
        expected_checksum: str = "",
        progress_callback: Optional[Callable[[Dict[str, Any]], None]] = None,
        cancellation_requested: Optional[Callable[[], bool]] = None,
    ) -> IngestionSummary:
        self.repository.initialize()
        self.repository.get_knowledge_base(knowledge_base_id)
        selected_count = (
            self.prepared_corpus_expected_documents
            if document_limit is None
            else max(1, min(int(document_limit), self.prepared_corpus_expected_documents))
        )
        _report_progress(
            progress_callback,
            "validation",
            2,
            current=0,
            total=selected_count,
        )
        documents, checksum = load_prepared_wixqa_corpus(
            self.prepared_corpus_path,
            expected_documents=self.prepared_corpus_expected_documents,
        )
        documents = documents[:selected_count]
        if expected_checksum and checksum != expected_checksum:
            raise KnowledgeProcessingError(
                "Prepared WixQA corpus changed after the import was queued. "
                "Refresh the source catalog and queue the import again."
            )
        _raise_if_cancelled(cancellation_requested)
        source_metadata = {
            "corpus_id": WIXQA_CORPUS_ID,
            "source": "Wix/WixQA",
            "source_revision": "main",
            "expected_document_count": self.prepared_corpus_expected_documents,
            "selected_document_count": selected_count,
            "selection_strategy": "corpus_order_prefix",
            "source_sha256": checksum,
            "source_path": str(self.prepared_corpus_path),
            "imported_at": utc_now(),
        }
        source = self.repository.create_data_source(
            knowledge_base_id=knowledge_base_id,
            source_type="wixqa_corpus",
            uri=WIXQA_CORPUS_URI,
            status="processing",
            metadata=source_metadata,
        )
        try:
            configuration = self._knowledge_base_configuration(knowledge_base_id)
            self.repository.update_knowledge_base(knowledge_base_id, status="processing", error=None)
            known_hashes = self.repository.existing_hashes(knowledge_base_id)
            stored: List[StoredKnowledgeDocument] = []
            skipped = 0
            for index, document in enumerate(documents, start=1):
                _raise_if_cancelled(cancellation_requested)
                text = normalize_structured_text(document.text)
                digest = content_hash(text)
                if digest in known_hashes:
                    skipped += 1
                    continue
                known_hashes.add(digest)
                stored.append(
                    StoredKnowledgeDocument(
                        id=f"doc-{uuid.uuid4().hex}",
                        knowledge_base_id=knowledge_base_id,
                        source_id=source.id,
                        title=document.title,
                        content_hash=digest,
                        text=text,
                        metadata={
                            **document.metadata,
                            **_processing_metadata(configuration),
                            "source_type": "wixqa_corpus",
                            "uri": document.uri,
                            "content_hash": digest,
                            "imported_at": utc_now(),
                        },
                    )
                )
                if index % 500 == 0:
                    _report_progress(
                        progress_callback,
                        "document_import",
                        5 + round(index * 20 / len(documents), 1),
                        current=index,
                        total=len(documents),
                    )
            self.repository.add_documents(stored)
            _report_progress(
                progress_callback,
                "document_import",
                25,
                current=len(documents),
                total=len(documents),
                documents_added=len(stored),
                documents_skipped=skipped,
            )
            reindex_summary = self.reindex(
                knowledge_base_id,
                progress_callback=progress_callback,
                cancellation_requested=cancellation_requested,
                progress_start=25,
            )
            refreshed = self.repository.get_knowledge_base(knowledge_base_id)
            index_details = self.query_embedding_details(knowledge_base_id)
            result_metadata = {
                **source_metadata,
                "documents_added": len(stored),
                "documents_skipped": skipped,
                "chunk_count": reindex_summary.chunks_added,
                "embedding_deployment_id": index_details.get("deployment_id", ""),
                "embedding_model": index_details.get("model", ""),
                "active_index_version_id": index_details.get("active_index_version_id", ""),
                "status": "completed",
                "completed_at": utc_now(),
            }
            self.repository.update_data_source(
                source.id,
                status="completed",
                metadata=result_metadata,
            )
            metadata = dict(refreshed.metadata)
            metadata["prepared_corpus"] = {
                "id": WIXQA_CORPUS_ID,
                "source": "Wix/WixQA",
                "expected_document_count": self.prepared_corpus_expected_documents,
                "selected_document_count": selected_count,
                "selection_strategy": "corpus_order_prefix",
                "imported_document_count": len(documents),
                "documents_added": len(stored),
                "documents_skipped": skipped,
                "chunk_count": reindex_summary.chunks_added,
                "embedding_deployment_id": index_details.get("deployment_id", ""),
                "embedding_model": index_details.get("model", ""),
                "active_index_version_id": index_details.get("active_index_version_id", ""),
                "sha256": checksum,
                "status": "completed",
                "indexed_at": utc_now(),
            }
            self.repository.update_knowledge_base(knowledge_base_id, metadata=metadata, error=None)
            summary = IngestionSummary(
                knowledge_base_id=knowledge_base_id,
                source_id=source.id,
                status=reindex_summary.status,
                documents_added=len(stored),
                documents_skipped=skipped,
                chunks_added=reindex_summary.chunks_added,
            )
            self.repository.record_ingestion_run(
                knowledge_base_id,
                summary.status,
                {
                    "documents_added": summary.documents_added,
                    "documents_skipped": summary.documents_skipped,
                    "chunks_added": summary.chunks_added,
                    "corpus_documents": len(documents),
                },
                source_id=source.id,
            )
            return summary
        except Exception as exc:
            terminal_status = "cancelled" if isinstance(exc, KnowledgeImportCancelled) else "failed"
            self.repository.update_data_source(
                source.id,
                status=terminal_status,
                metadata={
                    **source_metadata,
                    "status": terminal_status,
                    "error": str(exc),
                    "finished_at": utc_now(),
                },
            )
            current = self.repository.get_knowledge_base(knowledge_base_id)
            metadata = dict(current.metadata)
            metadata["prepared_corpus"] = {
                **dict(metadata.get("prepared_corpus") or {}),
                "id": WIXQA_CORPUS_ID,
                "source": "Wix/WixQA",
                "sha256": checksum,
                "status": terminal_status,
                "error": str(exc),
            }
            self.repository.update_knowledge_base(
                knowledge_base_id,
                status="indexed" if current.chunk_count else ("empty" if isinstance(exc, KnowledgeImportCancelled) else "failed"),
                metadata=metadata,
                error=None if isinstance(exc, KnowledgeImportCancelled) else str(exc),
            )
            raise

    def reindex(
        self,
        knowledge_base_id: str,
        *,
        progress_callback: Optional[Callable[[Dict[str, Any]], None]] = None,
        cancellation_requested: Optional[Callable[[], bool]] = None,
        progress_start: float = 0,
    ) -> IngestionSummary:
        self.repository.initialize()
        configuration = self._knowledge_base_configuration(knowledge_base_id)
        embedder = self._embedder_for_configuration(configuration)
        documents = self.repository.list_documents(knowledge_base_id)
        if configuration.get("chunking_strategy") == "structure_aware_recursive":
            documents = self._restore_wixqa_document_structure(documents)
        version = self._create_index_version(knowledge_base_id, configuration, embedder)
        try:
            self.repository.update_knowledge_base(knowledge_base_id, status="processing", error=None)
            chunks: List[StoredKnowledgeChunk] = []
            chunker = _chunker_from_configuration(configuration)
            for index, document in enumerate(documents, start=1):
                _raise_if_cancelled(cancellation_requested)
                chunks.extend(chunker.chunk_stored_document(_document_with_processing_metadata(document, configuration)))
                if index % 250 == 0:
                    _report_progress(
                        progress_callback,
                        "chunking",
                        progress_start + round(index * (20 if progress_start else 35) / max(len(documents), 1), 1),
                        current=index,
                        total=len(documents),
                        chunks_created=len(chunks),
                    )
            _report_progress(
                progress_callback,
                "chunking",
                progress_start + (20 if progress_start else 35),
                current=len(documents),
                total=len(documents),
                chunks_created=len(chunks),
            )
            if version is not None:
                chunks = [
                    replace(
                        chunk,
                        index_version_id=version.id,
                        metadata={**chunk.metadata, "index_version_id": version.id},
                    )
                    for chunk in chunks
                ]
            chunk_progress_start = progress_start + (20 if progress_start else 35)
            chunk_progress_span = 70 if progress_start else 60
            for batch_start in range(0, len(chunks), self.embedding_batch_size):
                _raise_if_cancelled(cancellation_requested)
                batch = chunks[batch_start : batch_start + self.embedding_batch_size]
                batch_end = batch_start + len(batch)
                batch_percent = min(
                    95,
                    chunk_progress_start
                    + round(batch_end * chunk_progress_span / max(len(chunks), 1), 1),
                )
                _report_progress(
                    progress_callback,
                    "embedding",
                    batch_percent,
                    current=batch_start,
                    total=len(chunks),
                    batch_size=len(batch),
                    embedding_model=embedder.model_name,
                )
                embedding_inputs = [
                    _embedding_input_for_chunk(chunk, configuration)
                    for chunk in batch
                ]
                embeddings = embedder.embed(embedding_inputs)
                _raise_if_cancelled(cancellation_requested)
                self.repository.append_chunks(batch, embeddings, embedder.model_name)
                _report_progress(
                    progress_callback,
                    "storage",
                    batch_percent,
                    current=batch_end,
                    total=len(chunks),
                    embedding_model=embedder.model_name,
                )
            _raise_if_cancelled(cancellation_requested)
            if version is not None:
                activate = getattr(self.repository, "activate_index_version", None)
                if callable(activate):
                    activate(knowledge_base_id, version.id)
            self.repository.update_knowledge_base(
                knowledge_base_id,
                status="indexed" if chunks else "empty",
                embedding_model=embedder.model_name,
                error=None,
            )
            _report_progress(
                progress_callback,
                "index_activation",
                100,
                current=len(chunks),
                total=len(chunks),
                index_version_id=version.id if version else "",
            )
        except Exception as exc:
            if version is not None:
                fail = getattr(self.repository, "fail_index_version", None)
                if callable(fail):
                    fail(knowledge_base_id, version.id, str(exc))
            current = self.repository.get_knowledge_base(knowledge_base_id)
            self.repository.update_knowledge_base(
                knowledge_base_id,
                status="indexed" if current.chunk_count else "failed",
                embedding_model=current.embedding_model,
                error=str(exc),
            )
            if isinstance(exc, KnowledgeProcessingError):
                raise
            raise KnowledgeProcessingError(str(exc)) from exc
        summary = IngestionSummary(
            knowledge_base_id=knowledge_base_id,
            source_id=None,
            status="indexed" if chunks else "empty",
            documents_added=0,
            documents_skipped=0,
            chunks_added=len(chunks),
        )
        self.repository.record_ingestion_run(
            knowledge_base_id,
            summary.status,
            {"documents_added": 0, "documents_skipped": 0, "chunks_added": len(chunks)},
        )
        return summary

    def _create_index_version(
        self,
        knowledge_base_id: str,
        configuration: Dict[str, Any],
        embedder: "EmbeddingModel",
    ) -> Optional[KnowledgeIndexVersionRecord]:
        create = getattr(self.repository, "create_index_version", None)
        if not callable(create):
            return None
        return create(
            knowledge_base_id,
            configuration,
            embedding_provider=configuration.get("embedding_provider", ""),
            embedding_model=embedder.model_name,
            embedding_deployment_id=configuration.get("embedding_deployment_id", ""),
            embedding_dimension=embedder.dimension,
        )

    def _restore_wixqa_document_structure(
        self,
        documents: List[StoredKnowledgeDocument],
    ) -> List[StoredKnowledgeDocument]:
        wixqa_documents = [
            document
            for document in documents
            if document.metadata.get("source_type") == "wixqa_corpus"
            or document.metadata.get("corpus_id") == WIXQA_CORPUS_ID
        ]
        if not wixqa_documents:
            return documents
        prepared, _ = load_prepared_wixqa_corpus(
            self.prepared_corpus_path,
            expected_documents=self.prepared_corpus_expected_documents,
        )
        source_text = {
            str(item.metadata.get("source_record_id") or ""): normalize_structured_text(item.text)
            for item in prepared
        }
        refreshed: List[StoredKnowledgeDocument] = []
        for document in documents:
            article_id = str(document.metadata.get("source_record_id") or "")
            structured_text = source_text.get(article_id)
            if structured_text and structured_text != document.text:
                document = replace(document, text=structured_text)
                self.repository.update_document(document)
            refreshed.append(document)
        return refreshed

    def _ingest_documents(
        self,
        knowledge_base_id: str,
        source_type: str,
        uri: str,
        documents: Iterable[KnowledgeDocumentInput],
        source_metadata: Dict[str, Any],
    ) -> IngestionSummary:
        source = self.repository.create_data_source(
            knowledge_base_id=knowledge_base_id,
            source_type=source_type,
            uri=uri,
            status="processing",
            metadata={**source_metadata, "imported_at": utc_now()},
        )
        try:
            configuration = self._knowledge_base_configuration(knowledge_base_id)
            self.repository.update_knowledge_base(knowledge_base_id, status="processing", error=None)
            known_hashes = self.repository.existing_hashes(knowledge_base_id)
            stored_documents: List[StoredKnowledgeDocument] = []
            skipped = 0
            for document in documents:
                text = normalize_structured_text(document.text)
                if not text:
                    skipped += 1
                    continue
                digest = content_hash(text)
                if digest in known_hashes:
                    skipped += 1
                    continue
                known_hashes.add(digest)
                stored_document = StoredKnowledgeDocument(
                    id=f"doc-{uuid.uuid4().hex}",
                    knowledge_base_id=knowledge_base_id,
                    source_id=source.id,
                    title=document.title or uri,
                    content_hash=digest,
                    text=text,
                    metadata={
                        **document.metadata,
                        **_processing_metadata(configuration),
                        "source_type": document.source_type,
                        "uri": document.uri,
                        "content_hash": digest,
                        "imported_at": utc_now(),
                    },
                )
                stored_documents.append(stored_document)
            self.repository.add_documents(stored_documents)
            reindex_summary = self.reindex(knowledge_base_id) if stored_documents else None
            chunks_added = reindex_summary.chunks_added if reindex_summary else 0
            status = reindex_summary.status if reindex_summary else ("indexed" if skipped else "empty")
            summary = IngestionSummary(
                knowledge_base_id=knowledge_base_id,
                source_id=source.id,
                status=status,
                documents_added=len(stored_documents),
                documents_skipped=skipped,
                chunks_added=chunks_added,
            )
            self.repository.record_ingestion_run(
                knowledge_base_id,
                status,
                {
                    "documents_added": len(stored_documents),
                    "documents_skipped": skipped,
                    "chunks_added": chunks_added,
                },
                source_id=source.id,
            )
            return summary
        except Exception as exc:
            message = str(exc)
            self.repository.update_knowledge_base(knowledge_base_id, status="failed", error=message)
            self.repository.record_ingestion_run(
                knowledge_base_id,
                "failed",
                {"documents_added": 0, "documents_skipped": 0, "chunks_added": 0},
                error=message,
                source_id=source.id,
            )
            if isinstance(exc, KnowledgeProcessingError):
                raise
            raise KnowledgeProcessingError(message) from exc

    def _knowledge_base_configuration(self, knowledge_base_id: str) -> Dict[str, Any]:
        knowledge_base = self.repository.get_knowledge_base(knowledge_base_id)
        return normalize_knowledge_base_configuration(knowledge_base.metadata.get("configuration"))

    def _active_index_configuration(self, knowledge_base_id: str) -> Dict[str, Any]:
        versions = self.list_index_versions(knowledge_base_id)
        active = next((version for version in versions if version.status == "active"), None)
        if active and active.chunking_configuration:
            return normalize_knowledge_base_configuration(active.chunking_configuration)
        return self._knowledge_base_configuration(knowledge_base_id)

    def _embedder_for_configuration(self, configuration: Dict[str, Any]) -> "EmbeddingModel":
        normalized = self._validate_configuration(configuration)
        deployment_id = normalized.get("embedding_deployment_id", "")
        if deployment_id and self.model_farm_service is not None and self.model_gateway is not None:
            deployment = self.model_farm_service.resolve(deployment_id, "embedding")
            cache_key = (deployment.provider, deployment.id, deployment.dimension)
            if cache_key not in self._embedder_cache:
                self._embedder_cache[cache_key] = GatewayEmbeddingModel(
                    self.model_gateway,
                    deployment.id,
                    model_name=deployment.model,
                    dimension=deployment.dimension,
                    external_processing_allowed=bool(normalized.get("external_processing_allowed")),
                )
            return self._embedder_cache[cache_key]
        provider = normalized["embedding_provider"]
        model_name = normalized["embedding_model"]
        cache_key = (provider, model_name, self._embedding_dimension)
        if cache_key not in self._embedder_cache:
            if model_name == HASH_EMBEDDING_MODEL:
                self._embedder_cache[cache_key] = HashEmbeddingModel(dimension=self._embedding_dimension)
            elif model_name == SENTENCE_TRANSFORMER_MINILM_MODEL:
                self._embedder_cache[cache_key] = SentenceTransformerEmbeddingModel(
                    model_name=model_name,
                    dimension=self._embedding_dimension,
                )
            else:
                supported = ", ".join(sorted(SUPPORTED_LOCAL_EMBEDDING_MODELS))
                raise KnowledgeProcessingError(
                    f"Embedding model {model_name!r} is not supported in v1. "
                    f"Modify the knowledge base to use Local with one of: {supported}."
                )
        return self._embedder_cache[cache_key]

    def _validate_configuration(self, configuration: Optional[Dict[str, Any]]) -> Dict[str, Any]:
        normalized = validate_knowledge_base_configuration(configuration)
        deployment_id = normalized.get("embedding_deployment_id", "")
        if deployment_id and self.model_farm_service is not None:
            try:
                deployment = self.model_farm_service.resolve(deployment_id, "embedding")
            except (KeyError, ModelFarmError) as exc:
                raise KnowledgeProcessingError(str(exc)) from exc
            if not deployment.is_local and not normalized.get("external_processing_allowed"):
                raise KnowledgeProcessingError(
                    "Enable external processing before selecting a remote embedding deployment."
                )
            normalized["embedding_provider"] = deployment.provider
            normalized["embedding_model"] = deployment.model
            if normalized["chunking_strategy"] == "structure_aware_recursive":
                if deployment.model != SENTENCE_TRANSFORMER_MINILM_MODEL or deployment.dimension != 384:
                    raise KnowledgeProcessingError(
                        "The WixQA structure-aware profile requires an enabled "
                        "sentence-transformers/all-MiniLM-L6-v2 embedding deployment with 384 dimensions."
                    )
        return normalized


class OverlapChunker:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 120):
        if chunk_size <= 0:
            raise ValueError("chunk_size must be positive")
        if chunk_overlap < 0 or chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be non-negative and smaller than chunk_size")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_stored_document(self, document: StoredKnowledgeDocument) -> List[StoredKnowledgeChunk]:
        text = clean_text(document.text)
        if not text:
            return []
        chunks: List[StoredKnowledgeChunk] = []
        start = 0
        index = 0
        step = self.chunk_size - self.chunk_overlap
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    StoredKnowledgeChunk(
                        id=f"chunk-{uuid.uuid4().hex}",
                        knowledge_base_id=document.knowledge_base_id,
                        document_id=document.id,
                        chunk_index=index,
                        text=chunk_text,
                        token_count=len(simple_tokens(chunk_text)),
                        metadata={
                            "title": document.title,
                            "source_id": document.source_id,
                            "content_hash": document.content_hash,
                            "start_char": start,
                            "end_char": end,
                            "chunking_mode": document.metadata.get("chunking_mode", "overlap"),
                            "chunking_strategy": document.metadata.get("chunking_strategy", document.metadata.get("chunking_mode", "overlap")),
                            "chunk_size": self.chunk_size,
                            "chunk_overlap": self.chunk_overlap,
                            "embedding_provider": document.metadata.get("embedding_provider", ""),
                            "embedding_model_requested": document.metadata.get("embedding_model_requested", ""),
                        },
                    )
                )
                index += 1
            if end >= len(text):
                break
            start += step
        return chunks


@dataclass(frozen=True)
class _StructuralUnit:
    text: str
    kind: str
    heading_path: Tuple[str, ...]
    section_id: str
    forced_split: bool = False


_WORDPIECE_TOKENIZER_CACHE: Dict[str, Any] = {}


def _load_wordpiece_tokenizer(model_name: str) -> Any:
    cached = _WORDPIECE_TOKENIZER_CACHE.get(model_name)
    if cached is not None:
        return cached
    try:
        from transformers import AutoTokenizer  # type: ignore

        tokenizer = AutoTokenizer.from_pretrained(model_name)
        backend = getattr(tokenizer, "backend_tokenizer", None)
        if backend is not None:
            backend.no_truncation()
            backend.no_padding()
    except ImportError as exc:
        raise KnowledgeProcessingError(
            "Install the ml extra to use structure-aware WordPiece chunking: "
            'python -m pip install -e ".[ml]".'
        ) from exc
    except Exception as exc:
        raise KnowledgeProcessingError(
            f"Unable to load the WordPiece tokenizer for {model_name!r}: {exc}. "
            'Install the ml extra and ensure the model is available locally.'
        ) from exc
    _WORDPIECE_TOKENIZER_CACHE[model_name] = tokenizer
    return tokenizer


class StructureAwareRecursiveChunker:
    def __init__(
        self,
        tokenizer: Any,
        *,
        options: Dict[str, Any],
        embedding_options: Dict[str, Any],
        profile_id: str = "",
    ):
        self.tokenizer = tokenizer
        self.options = dict(options)
        self.embedding_options = dict(embedding_options)
        self.profile_id = profile_id
        self.hard_max = int(self.embedding_options["hard_max_wordpieces"])
        self.target = int(self.options["target_body_tokens"])
        self.soft_max = int(self.options["soft_max_body_tokens"])
        self.overlap = int(self.options["overlap_tokens"])
        self.minimum = int(self.options["minimum_chunk_tokens"])
        self.separators = list(self.options["separators"])
        self.prefix_options = dict(self.options["metadata_prefix"])
        self.rules = dict(self.options["rules"])
        self._count_cache: Dict[Tuple[str, bool], int] = {}

    def _count(self, text: str, *, add_special_tokens: bool = False) -> int:
        key = (text, add_special_tokens)
        if key not in self._count_cache:
            self._count_cache[key] = _wordpiece_count(
                self.tokenizer,
                text,
                add_special_tokens=add_special_tokens,
            )
        return self._count_cache[key]

    def chunk_stored_document(self, document: StoredKnowledgeDocument) -> List[StoredKnowledgeChunk]:
        self._count_cache.clear()
        article_id = str(document.metadata.get("source_record_id") or document.id)
        units = self._parse_units(document, article_id)
        if not units:
            return []
        grouped: List[Tuple[_StructuralUnit, str, bool]] = []
        section_units: List[_StructuralUnit] = []
        current_section = ""
        for unit in units:
            if current_section and unit.section_id != current_section:
                grouped.extend(self._group_section(document, section_units))
                section_units = []
            current_section = unit.section_id
            section_units.append(unit)
        if section_units:
            grouped.extend(self._group_section(document, section_units))

        chunks: List[StoredKnowledgeChunk] = []
        previous_body = ""
        previous_section = ""
        for unit, body, forced_split in grouped:
            prefix = self._embedding_prefix(document, unit.heading_path)
            prefix_count = self._count(prefix)
            body_capacity = max(self.hard_max - prefix_count - _special_token_count(self.tokenizer), 1)
            overlap_text = ""
            if (
                chunks
                and self.overlap > 0
                and (
                    not self.rules.get("overlap_only_within_same_section", True)
                    or previous_section == unit.section_id
                )
            ):
                available = max(body_capacity - self._count(body), 0)
                overlap_text = _wordpiece_tail(
                    self.tokenizer,
                    previous_body,
                    min(self.overlap, available),
                )
            final_body = f"{overlap_text}\n\n{body}".strip() if overlap_text else body.strip()
            final_body = _truncate_wordpieces(self.tokenizer, final_body, body_capacity)
            body_count = self._count(final_body)
            overlap_count = self._count(overlap_text)
            embedding_text = _compose_embedding_text(prefix, final_body)
            embedding_count = self._count(embedding_text, add_special_tokens=True)
            while embedding_count > self.hard_max and body_count > 1:
                body_capacity = max(
                    body_count - max(embedding_count - self.hard_max, 1),
                    1,
                )
                final_body = _truncate_wordpieces(self.tokenizer, final_body, body_capacity)
                body_count = self._count(final_body)
                embedding_text = _compose_embedding_text(prefix, final_body)
                embedding_count = self._count(embedding_text, add_special_tokens=True)
            if embedding_count > self.hard_max:
                raise KnowledgeProcessingError(
                    "Unable to fit structure-aware chunk embedding input within "
                    f"{self.hard_max} WordPieces."
                )
            chunks.append(
                _stored_chunk(
                    document,
                    len(chunks),
                    final_body,
                    "structure_aware_recursive",
                    {
                        "article_id": article_id,
                        "parent_document_id": (
                            article_id
                            if self.options.get("parent_document") == "article_id"
                            else document.id
                        ),
                        "heading_path": list(unit.heading_path),
                        "section_id": unit.section_id,
                        "body_wordpiece_count": body_count,
                        "prefix_wordpiece_count": prefix_count,
                        "embedding_wordpiece_count": embedding_count,
                        "overlap_wordpiece_count": overlap_count,
                        "embedding_prefix": prefix,
                        "embedding_hard_max_wordpieces": self.hard_max,
                        "chunking_profile_id": self.profile_id,
                        "forced_split": bool(forced_split),
                        "below_minimum": body_count < self.minimum,
                    },
                    token_count=body_count,
                )
            )
            previous_body = body
            previous_section = unit.section_id
        return chunks

    def _parse_units(self, document: StoredKnowledgeDocument, article_id: str) -> List[_StructuralUnit]:
        text = normalize_structured_text(document.text)
        lines = text.splitlines()
        first_content = next((index for index, line in enumerate(lines) if line.strip()), None)
        if (
            first_content is not None
            and clean_text(lines[first_content]).casefold() == clean_text(document.title).casefold()
        ):
            remaining = lines[first_content + 1 :]
            lines = remaining if any(line.strip() for line in remaining) else [document.title]
        heading_path: List[str] = []
        units: List[_StructuralUnit] = []
        paragraph: List[str] = []
        list_block: List[str] = []

        def section_id() -> str:
            key = f"{article_id}|{' > '.join(heading_path)}"
            return hashlib.sha256(key.encode("utf-8")).hexdigest()[:16]

        def flush_paragraph() -> None:
            if paragraph:
                body = " ".join(item.strip() for item in paragraph if item.strip()).strip()
                if body:
                    units.append(
                        _StructuralUnit(body, "paragraph", tuple(heading_path), section_id())
                    )
                paragraph.clear()

        def flush_list() -> None:
            if list_block:
                body = "\n".join(item.strip() for item in list_block if item.strip()).strip()
                if body:
                    units.append(_StructuralUnit(body, "list", tuple(heading_path), section_id()))
                list_block.clear()

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                flush_list()
                flush_paragraph()
                continue
            heading = _heading_descriptor(line)
            heading_enabled = bool(
                heading is not None
                and (
                    (heading[0] <= 1 and "heading" in self.separators)
                    or (heading[0] > 1 and "subheading" in self.separators)
                )
            )
            if heading is not None and heading_enabled:
                flush_list()
                flush_paragraph()
                level, value = heading
                if level <= 1:
                    heading_path = [value]
                else:
                    heading_path = [*heading_path[: level - 1], value]
                continue
            is_numbered = bool(re.match(r"^\d+[.)]\s+", line))
            is_bullet = bool(re.match(r"^[-*+\u2022]\s+", line))
            preserve_list = (
                (is_numbered and self.rules.get("preserve_numbered_lists", True))
                or (is_bullet and self.rules.get("preserve_bullet_lists", True))
            )
            if preserve_list and "list" in self.separators:
                flush_paragraph()
                list_block.append(line)
                continue
            flush_list()
            paragraph.append(line)
        flush_list()
        flush_paragraph()
        return units

    def _group_section(
        self,
        document: StoredKnowledgeDocument,
        units: List[_StructuralUnit],
    ) -> List[Tuple[_StructuralUnit, str, bool]]:
        if not units:
            return []
        prefix = self._embedding_prefix(document, units[0].heading_path)
        capacity = max(
            self.hard_max
            - self._count(prefix)
            - _special_token_count(self.tokenizer),
            1,
        )
        soft_limit = min(self.soft_max, capacity)
        target = min(self.target, soft_limit)
        expanded: List[_StructuralUnit] = []
        for unit in units:
            expanded.extend(self._split_unit(unit, soft_limit))
        grouped: List[Tuple[_StructuralUnit, str, bool]] = []
        current: List[_StructuralUnit] = []
        for unit in expanded:
            candidate = _join_structural_units([*current, unit])
            candidate_count = self._count(candidate)
            if current and (
                candidate_count > soft_limit
                or self._count(_join_structural_units(current)) >= target
            ):
                grouped.append(
                    (
                        current[0],
                        _join_structural_units(current),
                        any(item.forced_split for item in current),
                    )
                )
                current = [unit]
            else:
                current.append(unit)
        if current:
            grouped.append(
                (
                    current[0],
                    _join_structural_units(current),
                    any(item.forced_split for item in current),
                )
            )
        if self.rules.get("merge_small_adjacent_chunks", True) and len(grouped) > 1:
            index = len(grouped) - 1
            while index > 0:
                unit, body, forced = grouped[index]
                if self._count(body) >= self.minimum:
                    index -= 1
                    continue
                previous_unit, previous_body, previous_forced = grouped[index - 1]
                merged = f"{previous_body}\n\n{body}".strip()
                if self._count(merged) <= soft_limit:
                    grouped[index - 1] = (
                        previous_unit,
                        merged,
                        previous_forced or forced,
                    )
                    grouped.pop(index)
                index -= 1
        return grouped

    def _split_unit(self, unit: _StructuralUnit, limit: int) -> List[_StructuralUnit]:
        if self._count(unit.text) <= limit:
            return [unit]
        parts = _recursive_wordpiece_split(
            unit.text,
            self.tokenizer,
            limit,
            self.separators,
            preserve_list=unit.kind == "list",
            count_tokens=self._count,
        )
        return [
            replace(unit, text=part, forced_split=True)
            for part in parts
            if part.strip()
        ]

    def _embedding_prefix(
        self,
        document: StoredKnowledgeDocument,
        heading_path: Tuple[str, ...],
    ) -> str:
        parts: List[str] = []
        if self.prefix_options.get("include_title", True) and document.title:
            parts.append(f"Title: {document.title}")
        if self.prefix_options.get("include_heading_path", True) and heading_path:
            parts.append(f"Section: {' > '.join(heading_path)}")
        if self.prefix_options.get("include_article_type", False):
            article_type = str(document.metadata.get("article_type") or "").strip()
            if article_type:
                parts.append(f"Article type: {article_type}")
        return _truncate_wordpieces(
            self.tokenizer,
            "\n".join(parts),
            int(self.prefix_options.get("maximum_tokens") or 0),
        )


class HeaderChunker:
    def __init__(self, chunk_size: int = 800):
        self.chunk_size = chunk_size

    def chunk_stored_document(self, document: StoredKnowledgeDocument) -> List[StoredKnowledgeChunk]:
        text = clean_text(document.text)
        sections: List[tuple[str, str]] = []
        heading = document.title
        body: List[str] = []
        for line in text.splitlines():
            clean = line.strip()
            if re.match(r"^#{1,6}\s+", clean) or (clean.endswith(":") and len(clean) < 100):
                if body:
                    sections.append((heading, "\n".join(body).strip()))
                heading = re.sub(r"^#{1,6}\s+", "", clean).rstrip(":").strip() or heading
                body = []
            else:
                body.append(line)
        if body:
            sections.append((heading, "\n".join(body).strip()))
        if not sections:
            sections = [(document.title, text)]
        chunks: List[StoredKnowledgeChunk] = []
        for section_heading, section_text in sections:
            for segment in _bounded_segments(section_text, self.chunk_size):
                chunks.append(_stored_chunk(document, len(chunks), segment, "header_based", {"header": section_heading}))
        return chunks


class RecursiveChunker:
    def __init__(self, chunk_size: int = 800, chunk_overlap: int = 0):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def chunk_stored_document(self, document: StoredKnowledgeDocument) -> List[StoredKnowledgeChunk]:
        segments = _recursive_split(clean_text(document.text), self.chunk_size, ["\n\n", "\n", ". ", " "])
        if self.chunk_overlap > 0 and len(segments) > 1:
            segments = [segments[0], *[(segments[index - 1][-self.chunk_overlap :] + " " + segment).strip() for index, segment in enumerate(segments[1:], 1)]]
        return [_stored_chunk(document, index, segment, "recursive") for index, segment in enumerate(segments) if segment]


class SemanticChunker:
    def __init__(self, chunk_size: int = 800, similarity_threshold: float = 0.12):
        self.chunk_size = chunk_size
        self.similarity_threshold = similarity_threshold

    def chunk_stored_document(self, document: StoredKnowledgeDocument) -> List[StoredKnowledgeChunk]:
        sentences = [item.strip() for item in re.split(r"(?<=[.!?])\s+|\n+", clean_text(document.text)) if item.strip()]
        if not sentences:
            return []
        groups: List[List[str]] = [[sentences[0]]]
        for sentence in sentences[1:]:
            current = " ".join(groups[-1])
            related = _token_similarity(current, sentence) >= self.similarity_threshold
            if related and len(current) + len(sentence) + 1 <= self.chunk_size:
                groups[-1].append(sentence)
            else:
                groups.append([sentence])
        segments: List[str] = []
        for group in groups:
            segments.extend(_bounded_segments(" ".join(group), self.chunk_size))
        return [
            _stored_chunk(document, index, segment, "semantic", {"semantic_threshold": self.similarity_threshold})
            for index, segment in enumerate(segments)
        ]


class HierarchicalChunker:
    def __init__(self, parent_size: int = 2000, child_size: int = 600, child_overlap: int = 80):
        self.parent_size = max(parent_size, child_size)
        self.child_size = child_size
        self.child_overlap = min(max(child_overlap, 0), child_size - 1)

    def chunk_stored_document(self, document: StoredKnowledgeDocument) -> List[StoredKnowledgeChunk]:
        parents = _bounded_segments(clean_text(document.text), self.parent_size)
        chunks: List[StoredKnowledgeChunk] = []
        for parent_index, parent_text in enumerate(parents):
            parent_id = f"parent-{uuid.uuid4().hex}"
            child_document = StoredKnowledgeDocument(
                document.id, document.knowledge_base_id, document.source_id, document.title,
                document.content_hash, parent_text, document.metadata,
            )
            children = OverlapChunker(self.child_size, self.child_overlap).chunk_stored_document(child_document)
            for child in children:
                chunks.append(
                    StoredKnowledgeChunk(
                        id=child.id,
                        knowledge_base_id=child.knowledge_base_id,
                        document_id=child.document_id,
                        chunk_index=len(chunks),
                        text=child.text,
                        token_count=child.token_count,
                        metadata={
                            **child.metadata,
                            "chunking_strategy": "hierarchical_parent_child",
                            "parent_chunk_id": parent_id,
                            "parent_chunk_index": parent_index,
                            "parent_text": parent_text,
                        },
                        parent_chunk_id=parent_id,
                    )
                )
        return chunks


def _stored_chunk(
    document: StoredKnowledgeDocument,
    index: int,
    text: str,
    strategy: str,
    extra_metadata: Optional[Dict[str, Any]] = None,
    token_count: Optional[int] = None,
) -> StoredKnowledgeChunk:
    return StoredKnowledgeChunk(
        id=f"chunk-{uuid.uuid4().hex}",
        knowledge_base_id=document.knowledge_base_id,
        document_id=document.id,
        chunk_index=index,
        text=text.strip(),
        token_count=token_count if token_count is not None else len(simple_tokens(text)),
        metadata={
            "title": document.title,
            "source_id": document.source_id,
            "content_hash": document.content_hash,
            "chunk_content_hash": content_hash(text),
            "chunking_mode": strategy,
            "chunking_strategy": strategy,
            "embedding_provider": document.metadata.get("embedding_provider", ""),
            "embedding_model_requested": document.metadata.get("embedding_model_requested", ""),
            "embedding_deployment_id": document.metadata.get("embedding_deployment_id", ""),
            **(extra_metadata or {}),
        },
    )


def _heading_descriptor(line: str) -> Optional[Tuple[int, str]]:
    markdown = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
    if markdown:
        return (1 if len(markdown.group(1)) == 1 else 2, markdown.group(2).strip())
    if line.endswith(":") and len(line) <= 120 and not re.match(r"^\w+://", line):
        return (2, line[:-1].strip())
    return None


def _join_structural_units(units: List[_StructuralUnit]) -> str:
    if not units:
        return ""
    output = units[0].text.strip()
    for previous, current in zip(units, units[1:]):
        separator = "\n" if previous.kind == "list" or current.kind == "list" else "\n\n"
        output = f"{output}{separator}{current.text.strip()}"
    return output.strip()


def _wordpiece_ids(tokenizer: Any, text: str, *, add_special_tokens: bool = False) -> List[int]:
    if not text:
        return []
    backend = getattr(tokenizer, "backend_tokenizer", None)
    if backend is not None:
        try:
            return list(
                backend.encode(
                    text,
                    add_special_tokens=add_special_tokens,
                ).ids
            )
        except (AttributeError, TypeError, ValueError):
            pass
    if callable(tokenizer):
        try:
            encoded = tokenizer(
                text,
                add_special_tokens=add_special_tokens,
                truncation=False,
                verbose=False,
            )
            if isinstance(encoded, dict) and "input_ids" in encoded:
                return list(encoded["input_ids"])
        except (TypeError, ValueError):
            pass
    return list(tokenizer.encode(text, add_special_tokens=add_special_tokens))


def _wordpiece_count(tokenizer: Any, text: str, *, add_special_tokens: bool = False) -> int:
    return len(_wordpiece_ids(tokenizer, text, add_special_tokens=add_special_tokens))


def _special_token_count(tokenizer: Any) -> int:
    try:
        return int(tokenizer.num_special_tokens_to_add(pair=False))
    except (AttributeError, TypeError, ValueError):
        return 2


def _decode_wordpieces(tokenizer: Any, token_ids: List[int]) -> str:
    if not token_ids:
        return ""
    return str(
        tokenizer.decode(
            token_ids,
            skip_special_tokens=True,
            clean_up_tokenization_spaces=True,
        )
    ).strip()


def _truncate_wordpieces(tokenizer: Any, text: str, maximum: int) -> str:
    if maximum <= 0 or not text:
        return ""
    token_ids = _wordpiece_ids(tokenizer, text)
    if len(token_ids) <= maximum:
        return text.strip()
    return _decode_wordpieces(tokenizer, token_ids[:maximum])


def _wordpiece_tail(tokenizer: Any, text: str, maximum: int) -> str:
    if maximum <= 0 or not text:
        return ""
    token_ids = _wordpiece_ids(tokenizer, text)
    return _decode_wordpieces(tokenizer, token_ids[-maximum:])


def _compose_embedding_text(prefix: str, body: str) -> str:
    return f"{prefix}\n\n{body}".strip() if prefix else body.strip()


def _embedding_input_for_chunk(
    chunk: StoredKnowledgeChunk,
    configuration: Dict[str, Any],
) -> str:
    if configuration.get("chunking_strategy") != "structure_aware_recursive":
        return chunk.text
    return _compose_embedding_text(
        str(chunk.metadata.get("embedding_prefix") or ""),
        chunk.text,
    )


def _recursive_wordpiece_split(
    text: str,
    tokenizer: Any,
    limit: int,
    separators: List[str],
    *,
    preserve_list: bool = False,
    count_tokens: Optional[Callable[[str], int]] = None,
) -> List[str]:
    clean = text.strip()
    count = count_tokens or (lambda value: _wordpiece_count(tokenizer, value))
    if not clean:
        return []
    if count(clean) <= limit:
        return [clean]
    for index, separator in enumerate(separators):
        parts: List[str] = []
        joiner = " "
        if separator == "paragraph":
            parts = [item.strip() for item in re.split(r"\n\s*\n", clean) if item.strip()]
            joiner = "\n\n"
        elif separator == "list" and (preserve_list or "\n" in clean):
            parts = [item.strip() for item in clean.splitlines() if item.strip()]
            joiner = "\n"
        elif separator == "sentence":
            parts = [
                item.strip()
                for item in re.split(r"(?<=[.!?])\s+", clean)
                if item.strip()
            ]
        elif separator == "token":
            token_ids = _wordpiece_ids(tokenizer, clean)
            return [
                _decode_wordpieces(tokenizer, token_ids[start : start + limit])
                for start in range(0, len(token_ids), limit)
            ]
        if len(parts) <= 1:
            continue
        fragments: List[str] = []
        remaining = separators[index + 1 :] or ["token"]
        for part in parts:
            fragments.extend(
                _recursive_wordpiece_split(
                    part,
                    tokenizer,
                    limit,
                    remaining,
                    preserve_list=preserve_list,
                    count_tokens=count,
                )
            )
        packed: List[str] = []
        buffer = ""
        for fragment in fragments:
            candidate = f"{buffer}{joiner if buffer else ''}{fragment}".strip()
            if buffer and count(candidate) > limit:
                packed.append(buffer)
                buffer = fragment
            else:
                buffer = candidate
        if buffer:
            packed.append(buffer)
        return packed
    token_ids = _wordpiece_ids(tokenizer, clean)
    return [
        _decode_wordpieces(tokenizer, token_ids[start : start + limit])
        for start in range(0, len(token_ids), limit)
    ]


def _bounded_segments(text: str, size: int) -> List[str]:
    clean = text.strip()
    if not clean:
        return []
    if len(clean) <= size:
        return [clean]
    return _recursive_split(clean, size, ["\n\n", "\n", ". ", " "])


def _recursive_split(text: str, size: int, separators: List[str]) -> List[str]:
    clean = text.strip()
    if not clean:
        return []
    if len(clean) <= size:
        return [clean]
    if not separators:
        return [clean[index : index + size].strip() for index in range(0, len(clean), size)]
    separator = separators[0]
    parts = clean.split(separator)
    if len(parts) == 1:
        return _recursive_split(clean, size, separators[1:])
    segments: List[str] = []
    buffer = ""
    for part in parts:
        candidate = f"{buffer}{separator if buffer else ''}{part}".strip()
        if len(candidate) <= size:
            buffer = candidate
            continue
        if buffer:
            segments.extend(_recursive_split(buffer, size, separators[1:]))
        buffer = part.strip()
    if buffer:
        segments.extend(_recursive_split(buffer, size, separators[1:]))
    return segments


def _token_similarity(left: str, right: str) -> float:
    left_tokens = set(simple_tokens(left))
    right_tokens = set(simple_tokens(right))
    return len(left_tokens & right_tokens) / max(len(left_tokens | right_tokens), 1)


class EmbeddingModel(Protocol):
    model_name: str
    dimension: int

    def embed(self, texts: List[str]) -> List[List[float]]:
        """Embed text chunks."""


class GatewayEmbeddingModel:
    def __init__(
        self,
        gateway: ModelGateway,
        deployment_id: str,
        *,
        model_name: str,
        dimension: int,
        external_processing_allowed: bool,
    ):
        self.gateway = gateway
        self.deployment_id = deployment_id
        self.model_name = model_name
        self.dimension = dimension
        self.external_processing_allowed = external_processing_allowed

    def embed(self, texts: List[str]) -> List[List[float]]:
        try:
            result = self.gateway.embed_sync(
                texts,
                self.deployment_id,
                context=ModelCallContext(purpose="knowledge_embedding"),
                external_processing_allowed=self.external_processing_allowed,
            )
        except ModelFarmError as exc:
            raise KnowledgeProcessingError(str(exc)) from exc
        self.dimension = result.dimension
        return result.embeddings


class HashEmbeddingModel:
    def __init__(self, dimension: int = 384, model_name: str = "hash-embedding-384"):
        self.dimension = dimension
        self.model_name = model_name

    def embed(self, texts: List[str]) -> List[List[float]]:
        return [self._embed_one(text) for text in texts]

    def _embed_one(self, text: str) -> List[float]:
        vector = [0.0] * self.dimension
        for token in simple_tokens(text):
            digest = hashlib.sha256(token.encode("utf-8")).hexdigest()
            vector[int(digest[:8], 16) % self.dimension] += 1.0
        norm = sum(value * value for value in vector) ** 0.5
        if norm == 0:
            return vector
        return [value / norm for value in vector]


class SentenceTransformerEmbeddingModel:
    def __init__(self, model_name: str, dimension: int = 384):
        self.model_name = model_name
        self.dimension = dimension
        self._tokenizer = None
        self._model = None
        self._torch = None

    def embed(self, texts: List[str]) -> List[List[float]]:
        if self._model is None:
            try:
                import torch  # type: ignore
                from transformers import AutoModel, AutoTokenizer  # type: ignore
            except ImportError as exc:
                raise KnowledgeProcessingError("Install the ml extra to use transformer embeddings.") from exc
            except Exception as exc:
                raise KnowledgeProcessingError(_sentence_transformer_runtime_error(self.model_name, exc)) from exc
            try:
                self._torch = torch
                self._tokenizer = AutoTokenizer.from_pretrained(self.model_name)
                self._model = AutoModel.from_pretrained(self.model_name)
                self._model.eval()
            except Exception as exc:
                raise KnowledgeProcessingError(_sentence_transformer_runtime_error(self.model_name, exc)) from exc
        try:
            encoded = self._tokenizer(
                texts,
                padding=True,
                truncation=True,
                max_length=512,
                return_tensors="pt",
            )
            with self._torch.no_grad():
                output = self._model(**encoded)
            token_embeddings = output.last_hidden_state
            attention_mask = encoded["attention_mask"].unsqueeze(-1).expand(token_embeddings.size()).float()
            pooled = (token_embeddings * attention_mask).sum(dim=1) / attention_mask.sum(dim=1).clamp(min=1e-9)
            normalized = self._torch.nn.functional.normalize(pooled, p=2, dim=1)
            embeddings = normalized.cpu().tolist()
        except Exception as exc:
            raise KnowledgeProcessingError(_sentence_transformer_runtime_error(self.model_name, exc)) from exc
        return [self._fit_dimension(list(map(float, row))) for row in embeddings]

    def _fit_dimension(self, vector: List[float]) -> List[float]:
        if len(vector) == self.dimension:
            return vector
        if len(vector) > self.dimension:
            fitted = vector[: self.dimension]
        else:
            fitted = [*vector, *([0.0] * (self.dimension - len(vector)))]
        norm = sum(value * value for value in fitted) ** 0.5
        if norm == 0:
            return fitted
        return [value / norm for value in fitted]


def build_embedder(model_name: str, dimension: int, use_sentence_transformers: bool) -> EmbeddingModel:
    if use_sentence_transformers:
        return SentenceTransformerEmbeddingModel(model_name=model_name, dimension=dimension)
    return HashEmbeddingModel(dimension=dimension)


def _sentence_transformer_runtime_error(model_name: str, exc: Exception) -> str:
    return (
        f"Unable to initialize transformer embedding model {model_name!r}: {exc}. "
        "Use hash-embedding-384 or repair the Windows CPU PyTorch dependency with "
        "python -m pip install --force-reinstall \"torch>=2.2,<2.6\" --index-url https://download.pytorch.org/whl/cpu "
        "then python -m pip install -e \".[dev,api,app,ml]\"."
    )


def normalize_knowledge_base_configuration(configuration: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    raw = dict(configuration or {})
    profile_id = str(raw.get("chunking_profile_id") or "").strip()
    profile_strategy = (
        "structure_aware_recursive"
        if profile_id == WIXQA_MINILM_PROFILE_ID
        else DEFAULT_KNOWLEDGE_BASE_CONFIGURATION["chunking_strategy"]
    )
    strategy = str(raw.get("chunking_strategy") or profile_strategy)
    if strategy not in CHUNKING_STRATEGIES:
        strategy = DEFAULT_KNOWLEDGE_BASE_CONFIGURATION["chunking_strategy"]
    chunk_size = _bounded_int(raw.get("chunk_size"), DEFAULT_KNOWLEDGE_BASE_CONFIGURATION["chunk_size"], minimum=100, maximum=12000)
    chunk_overlap = _bounded_int(raw.get("chunk_overlap"), DEFAULT_KNOWLEDGE_BASE_CONFIGURATION["chunk_overlap"], minimum=0, maximum=chunk_size - 1)
    if strategy == "fixed_size":
        chunk_overlap = 0
    provider = str(raw.get("embedding_provider") or DEFAULT_KNOWLEDGE_BASE_CONFIGURATION["embedding_provider"]).strip() or LOCAL_EMBEDDING_PROVIDER
    if provider.lower() == LOCAL_EMBEDDING_PROVIDER.lower():
        provider = LOCAL_EMBEDDING_PROVIDER
    profile_embedding_model = (
        SENTENCE_TRANSFORMER_MINILM_MODEL
        if profile_id == WIXQA_MINILM_PROFILE_ID
        else DEFAULT_KNOWLEDGE_BASE_CONFIGURATION["embedding_model"]
    )
    embedding_model = str(raw.get("embedding_model") or profile_embedding_model).strip() or HASH_EMBEDDING_MODEL
    embedding_deployment_id = str(raw.get("embedding_deployment_id") or "").strip()
    if (
        not embedding_deployment_id
        and provider == LOCAL_EMBEDDING_PROVIDER
        and embedding_model in SUPPORTED_LOCAL_EMBEDDING_MODELS
    ):
        embedding_deployment_id = (
            "model-local-minilm-384" if embedding_model == SENTENCE_TRANSFORMER_MINILM_MODEL else "model-local-hash-384"
        )
    raw_embedding_options = (
        dict(raw.get("embedding_options") or {})
        if isinstance(raw.get("embedding_options"), dict)
        else {}
    )
    hard_max_wordpieces = _bounded_int(
        raw_embedding_options.get("hard_max_wordpieces"),
        (
            WIXQA_MINILM_EMBEDDING_OPTIONS["hard_max_wordpieces"]
            if strategy == "structure_aware_recursive"
            else DEFAULT_EMBEDDING_OPTIONS["hard_max_wordpieces"]
        ),
        minimum=16,
        maximum=512,
    )
    raw_chunking_options = (
        dict(raw.get("chunking_options") or {})
        if isinstance(raw.get("chunking_options"), dict)
        else {}
    )
    raw_prefix = (
        dict(raw_chunking_options.get("metadata_prefix") or {})
        if isinstance(raw_chunking_options.get("metadata_prefix"), dict)
        else {}
    )
    raw_rules = (
        dict(raw_chunking_options.get("rules") or {})
        if isinstance(raw_chunking_options.get("rules"), dict)
        else {}
    )
    target_body_tokens = _bounded_int(
        raw_chunking_options.get("target_body_tokens"),
        WIXQA_MINILM_CHUNKING_OPTIONS["target_body_tokens"],
        minimum=16,
        maximum=max(hard_max_wordpieces - 2, 16),
    )
    soft_max_body_tokens = _bounded_int(
        raw_chunking_options.get("soft_max_body_tokens"),
        WIXQA_MINILM_CHUNKING_OPTIONS["soft_max_body_tokens"],
        minimum=target_body_tokens,
        maximum=max(hard_max_wordpieces - 2, target_body_tokens),
    )
    overlap_tokens = _bounded_int(
        raw_chunking_options.get("overlap_tokens"),
        WIXQA_MINILM_CHUNKING_OPTIONS["overlap_tokens"],
        minimum=0,
        maximum=max(target_body_tokens - 1, 0),
    )
    minimum_chunk_tokens = _bounded_int(
        raw_chunking_options.get("minimum_chunk_tokens"),
        WIXQA_MINILM_CHUNKING_OPTIONS["minimum_chunk_tokens"],
        minimum=1,
        maximum=target_body_tokens,
    )
    separators = []
    for separator in raw_chunking_options.get("separators") or WIXQA_MINILM_CHUNKING_OPTIONS["separators"]:
        name = str(separator).strip().lower()
        if name in STRUCTURE_SEPARATOR_TYPES and name not in separators:
            separators.append(name)
    if not separators:
        separators = list(WIXQA_MINILM_CHUNKING_OPTIONS["separators"])
    if "token" not in separators:
        separators.append("token")
    maximum_prefix_tokens = _bounded_int(
        raw_prefix.get("maximum_tokens"),
        WIXQA_MINILM_CHUNKING_OPTIONS["metadata_prefix"]["maximum_tokens"],
        minimum=0,
        maximum=max(hard_max_wordpieces - 3, 0),
    )
    chunking_options = {
        "parent_document": (
            str(raw_chunking_options.get("parent_document") or "article_id")
            if str(raw_chunking_options.get("parent_document") or "article_id") in {"article_id", "document_id"}
            else "article_id"
        ),
        "target_body_tokens": target_body_tokens,
        "soft_max_body_tokens": soft_max_body_tokens,
        "overlap_tokens": overlap_tokens,
        "minimum_chunk_tokens": minimum_chunk_tokens,
        "separators": separators,
        "metadata_prefix": {
            "include_title": bool(raw_prefix.get("include_title", True)),
            "include_heading_path": bool(raw_prefix.get("include_heading_path", True)),
            "include_article_type": bool(raw_prefix.get("include_article_type", False)),
            "maximum_tokens": maximum_prefix_tokens,
        },
        "rules": {
            key: (
                True
                if key == "never_merge_across_articles"
                and strategy == "structure_aware_recursive"
                else bool(raw_rules.get(key, default))
            )
            for key, default in WIXQA_MINILM_CHUNKING_OPTIONS["rules"].items()
        },
    }
    return {
        "chunking_strategy": strategy,
        "chunking_profile_id": profile_id,
        "chunk_size": chunk_size,
        "chunk_overlap": min(chunk_overlap, max(chunk_size - 1, 0)),
        "embedding_options": {"hard_max_wordpieces": hard_max_wordpieces},
        "chunking_options": chunking_options,
        "embedding_provider": provider,
        "embedding_model": embedding_model,
        "embedding_deployment_id": embedding_deployment_id,
        "external_processing_allowed": bool(raw.get("external_processing_allowed", False)),
    }


def validate_knowledge_base_configuration(configuration: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    normalized = normalize_knowledge_base_configuration(configuration)
    provider = normalized["embedding_provider"]
    model_name = normalized["embedding_model"]
    profile_id = normalized["chunking_profile_id"]
    if profile_id not in {"", WIXQA_MINILM_PROFILE_ID}:
        raise KnowledgeProcessingError(f"Unknown chunking profile {profile_id!r}.")
    if (
        profile_id == WIXQA_MINILM_PROFILE_ID
        and normalized["chunking_strategy"] != "structure_aware_recursive"
    ):
        raise KnowledgeProcessingError(
            f"Chunking profile {WIXQA_MINILM_PROFILE_ID!r} requires "
            "'structure_aware_recursive'."
        )
    if (
        normalized["chunking_strategy"] == "structure_aware_recursive"
        and model_name != SENTENCE_TRANSFORMER_MINILM_MODEL
    ):
        raise KnowledgeProcessingError(
            "Structure-aware recursive chunking requires "
            "'sentence-transformers/all-MiniLM-L6-v2' so WordPiece limits match the embedding model."
        )
    supported = ", ".join(sorted(SUPPORTED_LOCAL_EMBEDDING_MODELS))
    if normalized.get("embedding_deployment_id"):
        return normalized
    if provider != LOCAL_EMBEDDING_PROVIDER:
        raise KnowledgeProcessingError(
            f"Embedding provider {provider!r} is visible for the roadmap but is not executable in v1. "
            f"Modify the knowledge base to use Local with one of: {supported}."
        )
    if model_name not in SUPPORTED_LOCAL_EMBEDDING_MODELS:
        raise KnowledgeProcessingError(
            f"Embedding model {model_name!r} is not supported for Local execution in v1. "
            f"Modify the knowledge base to use one of: {supported}."
        )
    return normalized


def _bounded_int(value: Any, default: int, minimum: int, maximum: int) -> int:
    try:
        parsed = int(value)
    except (TypeError, ValueError):
        parsed = default
    return max(minimum, min(parsed, maximum))


def _chunker_from_configuration(configuration: Dict[str, Any]) -> Any:
    normalized = normalize_knowledge_base_configuration(configuration)
    strategy = normalized["chunking_strategy"]
    chunk_size = int(normalized["chunk_size"])
    overlap = int(normalized["chunk_overlap"])
    if strategy == "fixed_size":
        return OverlapChunker(chunk_size=chunk_size, chunk_overlap=0)
    if strategy == "sliding_window_overlap":
        return OverlapChunker(chunk_size=chunk_size, chunk_overlap=overlap)
    if strategy == "header_based":
        return HeaderChunker(chunk_size=chunk_size)
    if strategy == "recursive":
        return RecursiveChunker(chunk_size=chunk_size, chunk_overlap=overlap)
    if strategy == "structure_aware_recursive":
        return StructureAwareRecursiveChunker(
            _load_wordpiece_tokenizer(normalized["embedding_model"]),
            options=normalized["chunking_options"],
            embedding_options=normalized["embedding_options"],
            profile_id=normalized.get("chunking_profile_id", ""),
        )
    if strategy == "semantic":
        return SemanticChunker(chunk_size=chunk_size)
    if strategy == "hierarchical_parent_child":
        return HierarchicalChunker(
            parent_size=max(chunk_size * 3, 1200),
            child_size=chunk_size,
            child_overlap=overlap,
        )
    if strategy == "structure_aware_custom":
        raise KnowledgeProcessingError("Custom ara* chunking is disabled until the aratxt/arajson/aramd schemas are supplied.")
    raise KnowledgeProcessingError(f"Unsupported chunking strategy: {strategy}")


def _processing_metadata(configuration: Dict[str, Any]) -> Dict[str, Any]:
    normalized = normalize_knowledge_base_configuration(configuration)
    return {
        "chunking_strategy": normalized["chunking_strategy"],
        "chunking_profile_id": normalized["chunking_profile_id"],
        "chunking_mode": normalized["chunking_strategy"],
        "configured_chunk_size": normalized["chunk_size"],
        "configured_chunk_overlap": normalized["chunk_overlap"],
        "embedding_options": dict(normalized["embedding_options"]),
        "chunking_options": dict(normalized["chunking_options"]),
        "embedding_provider": normalized["embedding_provider"],
        "embedding_model_requested": normalized["embedding_model"],
        "embedding_deployment_id": normalized["embedding_deployment_id"],
        "external_processing_allowed": normalized["external_processing_allowed"],
    }


def _document_with_processing_metadata(
    document: StoredKnowledgeDocument,
    configuration: Dict[str, Any],
) -> StoredKnowledgeDocument:
    return StoredKnowledgeDocument(
        id=document.id,
        knowledge_base_id=document.knowledge_base_id,
        source_id=document.source_id,
        title=document.title,
        content_hash=document.content_hash,
        text=document.text,
        metadata={**document.metadata, **_processing_metadata(configuration)},
    )


def load_file_documents(filename: str, content: bytes) -> List[KnowledgeDocumentInput]:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_FILE_EXTENSIONS:
        raise KnowledgeProcessingError(f"Unsupported file type {extension!r}.")
    if extension in {".txt", ".md"}:
        return [_text_document(filename, content, "upload", extension)]
    if extension == ".json":
        return _json_documents(filename, content)
    if extension == ".jsonl":
        return _jsonl_documents(filename, content)
    if extension == ".pdf":
        return [_pdf_document(filename, content)]
    if extension == ".docx":
        return [_docx_document(filename, content)]
    if extension in CUSTOM_CHUNKING_EXTENSIONS:
        document = _text_document(filename, content, "upload", extension)
        return [
            KnowledgeDocumentInput(
                title=document.title,
                text=document.text,
                source_type=document.source_type,
                uri=document.uri,
                metadata={
                    **document.metadata,
                    "chunking_mode": "custom_placeholder",
                    "custom_loader_status": "schema_pending",
                },
            )
        ]
    raise KnowledgeProcessingError(f"No loader registered for {extension!r}.")


def load_prepared_wixqa_corpus(
    path: Path | str,
    *,
    expected_documents: int = WIXQA_CORPUS_EXPECTED_DOCUMENTS,
) -> Tuple[List[KnowledgeDocumentInput], str]:
    source_path = Path(path)
    if not source_path.is_file():
        raise KnowledgeProcessingError(
            "Prepared WixQA corpus is unavailable. Run "
            "`python scripts/download_wixqa.py --subset wixqa_expertwritten` first."
        )
    try:
        raw = source_path.read_bytes()
        decoded = raw.decode("utf-8")
    except (OSError, UnicodeDecodeError) as exc:
        raise KnowledgeProcessingError(f"Unable to read the prepared WixQA corpus: {exc}") from exc
    documents: List[KnowledgeDocumentInput] = []
    record_ids: set[str] = set()
    content_hashes: set[str] = set()
    try:
        for line_number, line in enumerate(decoded.splitlines(), start=1):
            if not line.strip():
                continue
            row = json.loads(line)
            if not isinstance(row, dict):
                raise KnowledgeProcessingError(f"WixQA corpus row {line_number} must be a JSON object.")
            record_id = str(row.get("id") or "").strip()
            text = str(row.get("text") or "").strip()
            metadata = row.get("metadata") or {}
            if not record_id or not text or not isinstance(metadata, dict):
                raise KnowledgeProcessingError(
                    f"WixQA corpus row {line_number} requires non-empty id/text and object metadata."
                )
            if record_id in record_ids:
                raise KnowledgeProcessingError(f"WixQA corpus contains duplicate record id {record_id!r}.")
            digest = content_hash(text)
            if digest in content_hashes:
                raise KnowledgeProcessingError(f"WixQA corpus contains duplicate content at row {line_number}.")
            record_ids.add(record_id)
            content_hashes.add(digest)
            first_line = next((part.strip() for part in text.splitlines() if part.strip()), record_id)
            documents.append(
                KnowledgeDocumentInput(
                    title=first_line[:240],
                    text=text,
                    source_type="wixqa_corpus",
                    uri=str(metadata.get("url") or f"{WIXQA_CORPUS_URI}/{record_id}"),
                    metadata={
                        **metadata,
                        "source": str(metadata.get("source") or "Wix/WixQA"),
                        "source_record_id": record_id,
                        "corpus_id": WIXQA_CORPUS_ID,
                        "corpus_revision": "main",
                    },
                )
            )
    except json.JSONDecodeError as exc:
        raise KnowledgeProcessingError(
            f"Prepared WixQA corpus contains invalid JSON at line {exc.lineno}: {exc.msg}."
        ) from exc
    if len(documents) != expected_documents:
        raise KnowledgeProcessingError(
            f"Prepared WixQA corpus must contain exactly {expected_documents:,} documents; "
            f"found {len(documents):,}."
        )
    return documents, hashlib.sha256(raw).hexdigest()


def _report_progress(
    callback: Optional[Callable[[Dict[str, Any]], None]],
    step: str,
    percent: float,
    **metadata: Any,
) -> None:
    if callback:
        callback({"step": step, "percent": round(float(percent), 1), **metadata})


def _raise_if_cancelled(callback: Optional[Callable[[], bool]]) -> None:
    if callback and callback():
        raise KnowledgeImportCancelled("WixQA corpus import was cancelled.")


def load_public_website(url: str) -> KnowledgeDocumentInput:
    _validate_public_url(url)
    request = urllib.request.Request(url, headers={"User-Agent": "aragbiz-ingestion/0.1"})
    with urllib.request.urlopen(request, timeout=20) as response:
        final_url = getattr(response, "geturl", lambda: url)()
        _validate_public_url(final_url)
        content_type = response.headers.get("content-type", "text/html").split(";", 1)[0].strip().lower()
        if content_type not in {"text/html", "application/xhtml+xml", "text/plain"}:
            raise KnowledgeProcessingError(f"Website content type {content_type!r} is not supported.")
        try:
            raw = response.read(5_000_001)
        except TypeError:
            raw = response.read()
        if len(raw) > 5_000_000:
            raise KnowledgeProcessingError("Website response exceeds the 5 MB ingestion limit.")
    html = raw.decode("utf-8", errors="replace")
    title = _html_title(html) or url
    text = _html_text(html)
    return KnowledgeDocumentInput(
        title=title,
        text=text,
        source_type="website",
        uri=final_url,
        metadata={
            "url": final_url,
            "content_type": content_type,
            "size_bytes": len(raw),
            "loader": "website_html",
        },
    )


def _validate_public_url(url: str) -> None:
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        raise KnowledgeProcessingError("Website source must be an absolute http:// or https:// URL.")
    hostname = parsed.hostname.lower().rstrip(".")
    if hostname in {"localhost", "localhost.localdomain"} or hostname.endswith((".local", ".internal")):
        raise KnowledgeProcessingError("Private or local website addresses are not allowed.")
    try:
        address = ipaddress.ip_address(hostname)
    except ValueError:
        return
    if not address.is_global:
        raise KnowledgeProcessingError("Private, loopback, link-local, and reserved website addresses are not allowed.")


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def normalize_structured_text(text: str) -> str:
    normalized = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t\f\v]+", " ", line).strip() for line in normalized.split("\n")]
    output: List[str] = []
    blank = False
    for line in lines:
        if line:
            output.append(line)
            blank = False
        elif output and not blank:
            output.append("")
            blank = True
    return "\n".join(output).strip()


def content_hash(text: str) -> str:
    return hashlib.sha256(clean_text(text).encode("utf-8")).hexdigest()


def simple_tokens(text: str) -> List[str]:
    return re.findall(r"[a-z0-9]+", text.lower())


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _text_document(filename: str, content: bytes, source_type: str, extension: str) -> KnowledgeDocumentInput:
    text = content.decode("utf-8", errors="replace")
    return KnowledgeDocumentInput(
        title=Path(filename).name,
        text=text,
        source_type=source_type,
        uri=filename,
        metadata=_file_metadata(filename, content, extension, loader="text"),
    )


def _json_documents(filename: str, content: bytes) -> List[KnowledgeDocumentInput]:
    payload = json.loads(content.decode("utf-8"))
    rows = payload.get("documents", payload) if isinstance(payload, dict) else payload
    if isinstance(rows, list):
        return [_document_from_json_row(filename, row, index) for index, row in enumerate(rows, start=1)]
    return [_document_from_json_row(filename, rows, 1)]


def _jsonl_documents(filename: str, content: bytes) -> List[KnowledgeDocumentInput]:
    documents: List[KnowledgeDocumentInput] = []
    for index, line in enumerate(content.decode("utf-8").splitlines(), start=1):
        if not line.strip():
            continue
        documents.append(_document_from_json_row(filename, json.loads(line), index))
    return documents


def _document_from_json_row(filename: str, row: Any, index: int) -> KnowledgeDocumentInput:
    if isinstance(row, dict):
        text = str(row.get("text") or row.get("content") or row.get("contents") or row.get("context") or json.dumps(row, ensure_ascii=True))
        title = str(row.get("title") or row.get("id") or f"{Path(filename).stem}-{index}")
        metadata = {key: value for key, value in row.items() if key not in {"text", "content", "contents", "context"}}
        if row.get("id") is not None:
            metadata["source_record_id"] = str(row["id"])
    else:
        text = str(row)
        title = f"{Path(filename).stem}-{index}"
        metadata = {}
    return KnowledgeDocumentInput(
        title=title,
        text=text,
        source_type="upload",
        uri=f"{filename}#{index}",
        metadata={**_file_metadata(filename, text.encode("utf-8"), Path(filename).suffix.lower(), loader="json"), **metadata},
    )


def _pdf_document(filename: str, content: bytes) -> KnowledgeDocumentInput:
    try:
        from pypdf import PdfReader  # type: ignore
    except ImportError as exc:
        raise KnowledgeProcessingError("Install the api extra to load PDF files.") from exc
    reader = PdfReader(BytesIO(content))
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    return KnowledgeDocumentInput(
        title=Path(filename).name,
        text=text,
        source_type="upload",
        uri=filename,
        metadata={**_file_metadata(filename, content, ".pdf", loader="pdf"), "page_count": len(reader.pages)},
    )


def _docx_document(filename: str, content: bytes) -> KnowledgeDocumentInput:
    try:
        from docx import Document as DocxDocument  # type: ignore
    except ImportError as exc:
        raise KnowledgeProcessingError("Install the api extra to load DOCX files.") from exc
    document = DocxDocument(BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return KnowledgeDocumentInput(
        title=Path(filename).name,
        text=text,
        source_type="upload",
        uri=filename,
        metadata={**_file_metadata(filename, content, ".docx", loader="docx"), "paragraph_count": len(document.paragraphs)},
    )


def _file_metadata(filename: str, content: bytes, extension: str, loader: str) -> Dict[str, Any]:
    return {
        "filename": Path(filename).name,
        "extension": extension,
        "mime_type": mimetypes.guess_type(filename)[0] or "application/octet-stream",
        "size_bytes": len(content),
        "loader": loader,
    }


def _html_text(html: str) -> str:
    try:
        from bs4 import BeautifulSoup  # type: ignore

        soup = BeautifulSoup(html, "html.parser")
        for node in soup(["script", "style", "noscript"]):
            node.decompose()
        return clean_text(soup.get_text(" "))
    except ImportError:
        parser = _TextExtractor()
        parser.feed(html)
        return clean_text(" ".join(parser.parts))


def _html_title(html: str) -> str:
    match = re.search(r"<title[^>]*>(.*?)</title>", html, flags=re.IGNORECASE | re.DOTALL)
    return clean_text(match.group(1)) if match else ""


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: List[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, Optional[str]]]) -> None:
        if tag in {"script", "style", "noscript"}:
            self._skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style", "noscript"} and self._skip_depth:
            self._skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self._skip_depth and data.strip():
            self.parts.append(data.strip())
